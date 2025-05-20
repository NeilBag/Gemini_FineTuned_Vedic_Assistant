import os
import logging
from flask import Flask, render_template, request, send_file, session, redirect, url_for, flash
from werkzeug.utils import secure_filename
from docx import Document
try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None
import google.generativeai as genai
from google.api_core import exceptions # For specific error handling
from dotenv import load_dotenv
import tempfile
import traceback

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__) # Use __name__ for logger

class Configuration:
    def __init__(self):
        self.DEBUG = os.getenv('FLASK_DEBUG', 'False').lower() in ('true', '1', 't')
        self.SECRET_KEY = os.getenv('FLASK_SECRET_KEY', os.urandom(24))
        self.UPLOAD_FOLDER = 'uploads'
        self.ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'txt', 'docx', 'mp4', 'avi', 'mov'}
        self.MAX_CONTENT_LENGTH = 16 * 1024 * 1024

class GenerativeAISettings:
    def __init__(self):
        self.MODEL_NAME = "tunedModels/hinduism-veda-expert-v1" # Your fine-tuned model
        # This is the prompt structure your Veda model was trained on
        self.PROMPT_TEMPLATE_VEDA = "Based on Hindu Vedic scriptures, answer the following question:\nQuestion: {user_question_and_context}\nAnswer:"
        self.GENERATION_CONFIG_DICT = { # Store as dict to pass to GenerationConfig object
            "temperature": 0.6,
            "top_p": 0.9,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }
        # self.CLIENT_OPTIONS = {"api_endpoint": "europe-west1-generativelanguage.googleapis.com"} # REMOVED - Will use global endpoint

# Global AI settings instance
ai_settings = GenerativeAISettings()

def create_app(config=Configuration()):
    app = Flask(__name__)
    app.config.from_object(config)

    api_key = os.getenv("GEMINI_API_KEY")
    if api_key:
        logger.info(f"DEBUG: GEMINI_API_KEY loaded. Length: {len(api_key)}. First 5: {api_key[:5]}, Last 5: {api_key[-5:]}")
        try:
            # !!!! CHANGED: Removed client_options !!!!
            genai.configure(
                api_key=api_key
            )
            logger.info("Generative AI SDK configured successfully (using default global endpoint).")
        except Exception as e:
            logger.exception(f"Failed to configure Generative AI SDK: {e}")
    else:
        logger.error("CRITICAL: GEMINI_API_KEY is NOT SET in the environment. AI features will not work.")

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    return app

app = create_app()

# Initialize the AI model globally
model = None
if os.getenv("GEMINI_API_KEY"): # Check if API key was set for configuration
    try:
        model = genai.GenerativeModel(
            model_name=ai_settings.MODEL_NAME
        )
        logger.info(f"Global GenerativeModel '{ai_settings.MODEL_NAME}' initialized (using default global endpoint).")
    except Exception as e:
        logger.exception(f"Failed to initialize global GenerativeModel '{ai_settings.MODEL_NAME}': {e}")
else:
    logger.warning("Global AI Model not initialized because GEMINI_API_KEY is not set or SDK configuration failed.")


# --- File Handling ---
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def secure_save_file(file):
    if not file or not file.filename: return None, None
    filename = secure_filename(file.filename)
    if not filename: filename = "unnamed_file_" + tempfile.mktemp().split(os.sep)[-1]
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    try:
        file.save(file_path)
        logger.info(f"File saved: {file_path}")
        return file_path, file.filename
    except Exception as e:
        logger.exception(f"Failed to save file {filename}: {e}")
        return None, None

def read_file_content(file_path):
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    content = ""
    try:
        if extension == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f: content = f.read()
        elif extension == '.docx':
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
        else:
            logger.warning(f"Content reading not supported for extension: {extension}")
            return f"(File type {extension} not processed for content)"
        max_chars = 15000
        if len(content) > max_chars:
            content = content[:max_chars] + "\n... [truncated]"
        return content
    except Exception as e:
        logger.exception(f"Error reading file content from {file_path}: {e}")
        return f"(Error reading file {os.path.basename(file_path)})"

# --- History Handling ---
def build_history_for_template(session_history):
    template_history = []
    for entry in session_history:
        template_history.append({'user': entry.get('user'), 'model': entry.get('model')})
    return template_history

# --- Export Functions ---
def export_docx(text):
    try:
        document = Document()
        document.add_paragraph(text)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            docx_path = temp_file.name
            document.save(docx_path)
        logger.info(f"Generated DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        logger.exception(f"Error creating DOCX file: {e}")
        return None

def export_pdf(text):
    if pisa is None:
        logger.error("xhtml2pdf library is not available. Cannot export PDF.")
        return None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            pdf_path = temp_file.name
            html_content = f"<!DOCTYPE html><html><head><meta charset='UTF-8'></head><body><pre>{text}</pre></body></html>"
            from io import BytesIO
            pdf_buffer = BytesIO()
            pisa_status = pisa.CreatePDF(BytesIO(html_content.encode('UTF-8')), dest=pdf_buffer)
            if pisa_status.err:
                 logger.error(f"Error creating PDF using pisa: {pisa_status.err}")
                 return None
            with open(pdf_path, "wb") as f: f.write(pdf_buffer.getvalue())
        logger.info(f"Generated PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        logger.exception(f"Error creating PDF file: {e}")
        return None

# --- Routes ---
@app.route("/", methods=["GET", "POST"])
def index2():
    if 'chat_history' not in session:
        session['chat_history'] = []

    if model is None:
        flash("AI Model is not configured or failed to initialize. Please check the GEMINI_API_KEY and server logs.", "error")
        template_history = build_history_for_template(session['chat_history'])
        return render_template("index2.html", history=template_history)

    generation_config_obj = genai.types.GenerationConfig(**ai_settings.GENERATION_CONFIG_DICT)
    user_prompt_for_history = ""
    full_content_for_model = ""

    if request.method == "POST":
        prompt_for_api = ""
        try:
            user_text_prompt = request.form.get("prompt", "").strip()
            user_prompt_for_history = user_text_prompt
            export_format = request.form.get("export_format", "")

            uploaded_file_contents_str = ""
            uploaded_filenames = []
            files = request.files.getlist("files")
            if files:
                temp_file_contents = []
                for file_in in files:
                    if file_in and file_in.filename and allowed_file(file_in.filename):
                        file_path, original_filename = secure_save_file(file_in)
                        if file_path:
                            logger.info(f"Processing uploaded file: {original_filename}")
                            content = read_file_content(file_path)
                            temp_file_contents.append(f"--- Context from {original_filename} ---\n{content}\n--- End of {original_filename} Context ---")
                            uploaded_filenames.append(original_filename)
                        else:
                            flash(f"Could not save file: {file_in.filename}", "warning")
                    elif file_in and file_in.filename:
                        flash(f"File type not allowed: {file_in.filename}", "warning")
                if temp_file_contents:
                    uploaded_file_contents_str = "\n\n".join(temp_file_contents)

            if uploaded_file_contents_str:
                full_content_for_model = uploaded_file_contents_str + "\n\nUser Question: " + user_text_prompt
                if uploaded_filenames:
                     user_prompt_for_history += f"\n\n(Used uploaded files: {', '.join(uploaded_filenames)})"
            else:
                full_content_for_model = user_text_prompt

            if not full_content_for_model.strip():
                flash("Please enter a prompt or upload a file with content.", "warning")
                template_history = build_history_for_template(session['chat_history'])
                return render_template("index2.html", history=template_history)

            prompt_for_api = ai_settings.PROMPT_TEMPLATE_VEDA.format(user_question_and_context=full_content_for_model)
            logger.info("Sending prompt to Generative AI using generate_content...")
            logger.info(f"Formatted prompt for API (first 300 chars): {prompt_for_api[:300]}...")

            response = model.generate_content(
                prompt_for_api,
                generation_config=generation_config_obj,
            )
            response_text = response.text
            logger.info(f"Received response from Generative AI (first 300 chars): {response_text[:300]}...")

            session['chat_history'].append({'user': user_prompt_for_history, 'model': response_text})
            session.modified = True

            if export_format:
                export_path, mimetype, download_name = None, None, None
                if export_format == "docx":
                    export_path = export_docx(response_text)
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    download_name="response.docx"
                elif export_format == "pdf":
                    if pisa:
                        export_path = export_pdf(response_text)
                        mimetype='application/pdf'
                        download_name="output.pdf"
                    else: flash("PDF export is currently unavailable.", "error")

                if export_path and mimetype and download_name:
                    try: return send_file(export_path, mimetype=mimetype, as_attachment=True, download_name=download_name)
                    finally:
                        try: os.remove(export_path); logger.info(f"Cleaned up: {export_path}")
                        except OSError as e: logger.error(f"Error removing temp file {export_path}: {e}")
                elif not export_path and export_format == "pdf" and not pisa: pass
                else: flash(f"Failed to export as {export_format}.", "error")

            template_history = build_history_for_template(session['chat_history'])
            return render_template("index2.html", history=template_history)

        except exceptions.InvalidArgument as e:
            logger.error(f"Problematic full_content_for_model (first 300 chars): {full_content_for_model[:300]}")
            logger.error(f"Problematic prompt_for_api (first 300 chars): {prompt_for_api[:300] if 'prompt_for_api' in locals() else 'Not generated'}")
            logger.error(f"InvalidArgument error processing POST request: {str(e)}")
            logger.error(traceback.format_exc())
            flash(f"An API error occurred (InvalidArgument): {str(e)}. This often means the prompt structure is not what the fine-tuned model expects or the content is problematic.", "error")
            return redirect(url_for('index2'))
        except exceptions.MethodNotImplemented as e: # Catch the specific error
            logger.error(f"Problematic full_content_for_model (first 300 chars): {full_content_for_model[:300]}")
            logger.error(f"Problematic prompt_for_api (first 300 chars): {prompt_for_api[:300] if 'prompt_for_api' in locals() else 'Not generated'}")
            logger.error(f"MethodNotImplemented error (501/404) processing POST request: {str(e)}")
            logger.error(traceback.format_exc())
            flash(f"An API error occurred (MethodNotImplemented/404): {str(e)}. This usually means the model name is incorrect, or the API endpoint cannot find/serve this specific fine-tuned model. Try removing regional client_options if set.", "error")
            return redirect(url_for('index2'))
        except Exception as e:
            logger.error(f"Problematic full_content_for_model (first 300 chars): {full_content_for_model[:300]}")
            logger.error(f"General error processing POST request: {str(e)}")
            logger.error(traceback.format_exc())
            flash(f"An error occurred: {str(e)}", "error")
            return redirect(url_for('index2'))

    template_history = build_history_for_template(session['chat_history'])
    return render_template("index2.html", history=template_history)

@app.route("/clear", methods=["POST"])
def clear_history():
    session.pop('chat_history', None)
    flash("Chat history cleared.", "info")
    return redirect(url_for('index2'))

@app.errorhandler(413)
def request_entity_too_large(e):
    flash("The uploaded file(s) are too large. Maximum allowed total size is 16MB.", "error")
    return redirect(url_for('index2'))

@app.route("/minimal_test")
def minimal_test_route():
    logger = logging.getLogger(__name__)
    logger.info("Executing /minimal_test route (using default global endpoint)")

    if not os.getenv("GEMINI_API_KEY"):
        logger.error("MINIMAL_TEST: GEMINI_API_KEY not found.")
        return "Error: GEMINI_API_KEY not found", 500

    if model is None:
        logger.error("MINIMAL_TEST: Global model is not initialized.")
        return "Error: Global model not initialized. Check server logs.", 500

    test_model_name = ai_settings.MODEL_NAME
    try:
        logger.info(f"MINIMAL_TEST: Attempting to generate content with global model: {model.model_name}")

        generation_config_test = genai.types.GenerationConfig(
            temperature=0.7, max_output_tokens=50
        )
        test_prompt_minimal = "Briefly explain the concept of Atman in Vedanta." # Veda related

        # IMPORTANT: Format the prompt for the Q&A model
        formatted_test_prompt = ai_settings.PROMPT_TEMPLATE_VEDA.format(user_question_and_context=test_prompt_minimal)
        logger.info(f"MINIMAL_TEST: Formatted prompt: {formatted_test_prompt}")

        response = model.generate_content( # Use the global 'model' instance
            formatted_test_prompt,
            generation_config=generation_config_test
            )

        response_text = response.text if hasattr(response, 'text') and response.text else str(response)
        logger.info(f"MINIMAL_TEST: Response received: {response_text}")
        return f"Minimal test successful. Response: {response_text}", 200
    except exceptions.MethodNotImplemented as e: # Catch the specific error
        logger.error(f"MINIMAL_TEST: MethodNotImplemented error (501/404) for model {test_model_name}: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error during minimal test (MethodNotImplemented/404): {str(e)}. This means the model couldn't be found or served via the default global endpoint. Check model name and its status in AI Studio/Vertex AI.", 500
    except Exception as e:
        logger.error(f"MINIMAL_TEST: Error during minimal test for model {test_model_name}: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error during minimal test: {str(e)}", 500

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', 5000)), debug=app.config['DEBUG'])
